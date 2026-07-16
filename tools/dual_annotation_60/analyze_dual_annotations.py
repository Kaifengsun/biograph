"""Analyze independent Reviewer A/B labels and prepare a compact disagreement registry."""

from __future__ import annotations

import argparse
import json
import math
import random
import statistics
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[2]
CORE_LABELS = {"Sufficient", "Required Component"}
PASSAGE_LABELS = ["Sufficient", "Required Component", "Context Only", "Not Supporting", "Unclear"]
QUESTION_LABELS = ["Answerable", "Needs Revision", "Invalid"]
COMPLETENESS_LABELS = ["Complete", "Incomplete", "Unclear"]
SEED = 20260716


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def parse_a_docx(path: Path) -> tuple[dict[str, dict[str, str]], dict[tuple[str, str], dict[str, str]]]:
    document = Document(path)
    if len(document.tables) < 44:
        raise ValueError(f"unexpected Reviewer A DOCX structure: {len(document.tables)} tables")
    questions: dict[str, dict[str, str]] = {}
    passages: dict[tuple[str, str], dict[str, str]] = {}

    for table_index in (0, 21, 42):
        table = document.tables[table_index]
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            values = [cell.text.strip() for cell in row.cells]
            record = dict(zip(headers, values, strict=False))
            query_id = record.get("Review ID", "")
            if not query_id:
                continue
            questions[query_id] = {
                "question_status": record.get("Question Status", record.get("Question Status*", "")),
                "evidence_set_completeness": record.get("Evidence Set Completeness", record.get("Evidence Set Completeness*", "")),
                "reviewer_note": record.get("Reviewer Note", record.get("简要理由", "")),
            }

    batch_01_ids = [row.cells[1].text.strip() for row in document.tables[0].rows[1:]]
    for table_index, query_id in enumerate(batch_01_ids, 1):
        for row in document.tables[table_index].rows[1:]:
            passage_id, label = row.cells[0].text.strip(), row.cells[1].text.strip()
            passages[(query_id, passage_id)] = {"passage_label": label, "reviewer_note": ""}

    for table_index in range(22, 42):
        for row in document.tables[table_index].rows[1:]:
            query_id, passage_id, label = [cell.text.strip() for cell in row.cells[:3]]
            note = row.cells[3].text.strip() if len(row.cells) > 3 else ""
            passages[(query_id, passage_id)] = {"passage_label": label, "reviewer_note": note}

    for row in document.tables[43].rows[1:]:
        query_id, passage_id, label = [cell.text.strip() for cell in row.cells[:3]]
        passages[(query_id, passage_id)] = {"passage_label": label, "reviewer_note": ""}
    return questions, passages


def parse_b_workbooks(paths: Iterable[Path]) -> tuple[dict[str, dict[str, str]], dict[tuple[str, str], dict[str, str]]]:
    questions: dict[str, dict[str, str]] = {}
    passages: dict[tuple[str, str], dict[str, str]] = {}
    for path in sorted(paths):
        workbook = load_workbook(path, data_only=False)
        if len(workbook.worksheets) != 4:
            raise ValueError(f"unexpected Reviewer B workbook structure: {path}")
        question_sheet, evidence_sheet = workbook.worksheets[1], workbook.worksheets[2]
        for row in range(5, question_sheet.max_row + 1):
            query_id = str(question_sheet.cell(row, 2).value or "").strip()
            if not query_id:
                continue
            if query_id in questions:
                raise ValueError(f"duplicate Reviewer B query ID: {query_id}")
            questions[query_id] = {
                "question_status": str(question_sheet.cell(row, 5).value or "").strip(),
                "evidence_set_completeness": str(question_sheet.cell(row, 6).value or "").strip(),
                "reviewer_note": str(question_sheet.cell(row, 7).value or "").strip(),
            }
        for row in range(5, evidence_sheet.max_row + 1):
            query_id = str(evidence_sheet.cell(row, 2).value or "").strip()
            passage_id = str(evidence_sheet.cell(row, 5).value or "").strip()
            if not query_id or not passage_id:
                continue
            key = (query_id, passage_id)
            if key in passages:
                raise ValueError(f"duplicate Reviewer B passage: {key}")
            passages[key] = {
                "passage_label": str(evidence_sheet.cell(row, 9).value or "").strip(),
                "reviewer_note": str(evidence_sheet.cell(row, 10).value or "").strip(),
            }
    return questions, passages


def validate_labels(
    questions: dict[str, dict[str, str]],
    passages: dict[tuple[str, str], dict[str, str]],
    reviewer: str,
) -> list[dict[str, str]]:
    issues = []
    for query_id, record in questions.items():
        if record["question_status"] not in QUESTION_LABELS:
            issues.append({"reviewer": reviewer, "key": query_id, "field": "question_status", "value": record["question_status"]})
        if record["evidence_set_completeness"] not in COMPLETENESS_LABELS:
            issues.append({"reviewer": reviewer, "key": query_id, "field": "evidence_set_completeness", "value": record["evidence_set_completeness"]})
    for key, record in passages.items():
        if record["passage_label"] not in PASSAGE_LABELS:
            issues.append({"reviewer": reviewer, "key": "/".join(key), "field": "passage_label", "value": record["passage_label"]})
    return issues


def confusion(labels_a: list[str], labels_b: list[str], categories: list[str]) -> dict[str, dict[str, int]]:
    return {a: {b: sum(x == a and y == b for x, y in zip(labels_a, labels_b, strict=True)) for b in categories} for a in categories}


def cohen_kappa(labels_a: list[str], labels_b: list[str], categories: list[str]) -> float | None:
    if not labels_a:
        return None
    n = len(labels_a)
    observed = sum(a == b for a, b in zip(labels_a, labels_b, strict=True)) / n
    expected = sum((labels_a.count(category) / n) * (labels_b.count(category) / n) for category in categories)
    return None if math.isclose(expected, 1.0) else (observed - expected) / (1.0 - expected)


def gwet_ac1(labels_a: list[str], labels_b: list[str], categories: list[str]) -> float | None:
    if not labels_a:
        return None
    n = len(labels_a)
    observed = sum(a == b for a, b in zip(labels_a, labels_b, strict=True)) / n
    marginals = [((labels_a.count(category) + labels_b.count(category)) / (2 * n)) for category in categories]
    expected = sum(p * (1 - p) for p in marginals) / max(1, len(categories) - 1)
    return None if math.isclose(expected, 1.0) else (observed - expected) / (1.0 - expected)


def agreement_metrics(labels_a: list[str], labels_b: list[str], categories: list[str]) -> dict[str, Any]:
    n = len(labels_a)
    return {
        "n": n,
        "exact_agreement": None if not n else sum(a == b for a, b in zip(labels_a, labels_b, strict=True)) / n,
        "cohen_kappa": cohen_kappa(labels_a, labels_b, categories),
        "gwet_ac1": gwet_ac1(labels_a, labels_b, categories),
        "confusion_matrix": confusion(labels_a, labels_b, categories),
    }


def binary_metrics(labels_a: list[bool], labels_b: list[bool]) -> dict[str, Any]:
    a = sum(x and y for x, y in zip(labels_a, labels_b, strict=True))
    b = sum(x and not y for x, y in zip(labels_a, labels_b, strict=True))
    c = sum(not x and y for x, y in zip(labels_a, labels_b, strict=True))
    d = sum(not x and not y for x, y in zip(labels_a, labels_b, strict=True))
    text_a = ["Gold" if value else "Non-Gold" for value in labels_a]
    text_b = ["Gold" if value else "Non-Gold" for value in labels_b]
    result = agreement_metrics(text_a, text_b, ["Gold", "Non-Gold"])
    result.update({
        "both_gold": a,
        "a_gold_b_non_gold": b,
        "a_non_gold_b_gold": c,
        "both_non_gold": d,
        "positive_agreement": None if 2 * a + b + c == 0 else 2 * a / (2 * a + b + c),
        "negative_agreement": None if 2 * d + b + c == 0 else 2 * d / (2 * d + b + c),
    })
    return result


def percentile(values: list[float], probability: float) -> float:
    ordered = sorted(values)
    index = min(len(ordered) - 1, max(0, int(round((len(ordered) - 1) * probability))))
    return ordered[index]


def cluster_bootstrap_core(
    paired_by_query: dict[str, list[tuple[bool, bool]]],
    iterations: int,
    seed: int,
) -> dict[str, dict[str, float]]:
    rng = random.Random(seed)
    query_ids = sorted(paired_by_query)
    samples: dict[str, list[float]] = defaultdict(list)
    for _ in range(iterations):
        pairs = []
        for _index in range(len(query_ids)):
            pairs.extend(paired_by_query[rng.choice(query_ids)])
        metrics = binary_metrics([a for a, _b in pairs], [b for _a, b in pairs])
        for key in ("exact_agreement", "cohen_kappa", "gwet_ac1", "positive_agreement"):
            value = metrics[key]
            if value is not None:
                samples[key].append(float(value))
    return {
        key: {"ci_95_low": percentile(values, 0.025), "ci_95_high": percentile(values, 0.975)}
        for key, values in samples.items() if values
    }


def per_query_set_metrics(
    passage_a: dict[tuple[str, str], dict[str, str]],
    passage_b: dict[tuple[str, str], dict[str, str]],
    query_ids: Iterable[str],
) -> list[dict[str, Any]]:
    rows = []
    for query_id in sorted(query_ids):
        gold_a = {passage_id for (qid, passage_id), record in passage_a.items() if qid == query_id and record["passage_label"] in CORE_LABELS}
        gold_b = {passage_id for (qid, passage_id), record in passage_b.items() if qid == query_id and record["passage_label"] in CORE_LABELS}
        union = gold_a | gold_b
        intersection = gold_a & gold_b
        if not union:
            jaccard = f1 = None
        else:
            jaccard = len(intersection) / len(union)
            f1 = 2 * len(intersection) / (len(gold_a) + len(gold_b)) if gold_a or gold_b else None
        rows.append({
            "query_id": query_id,
            "gold_count_a": len(gold_a),
            "gold_count_b": len(gold_b),
            "intersection": len(intersection),
            "union": len(union),
            "jaccard": jaccard,
            "f1": f1,
        })
    return rows


def markdown_report(report: dict[str, Any]) -> str:
    question = report["agreement"]["question_status"]
    completeness = report["agreement"]["evidence_set_completeness"]
    core = report["agreement"]["passage_core_binary"]
    five = report["agreement"]["passage_five_way"]
    set_metrics = report["per_query_gold_set_summary"]
    lines = [
        "# Independent Dual-Annotation Agreement Report",
        "",
        f"- Reviewer A question labels: {report['counts']['questions_a']}/60",
        f"- Reviewer B question labels: {report['counts']['questions_b']}/60",
        f"- Paired passage labels: {report['counts']['paired_passages']}/529",
        f"- Missing passage labels: A={report['counts']['missing_passages_a']}, B={report['counts']['missing_passages_b']}",
        "",
        "## Agreement",
        "",
        "| Level | Exact agreement | Cohen's kappa | Gwet AC1 |",
        "|---|---:|---:|---:|",
        f"| Question status | {question['exact_agreement']:.3f} | {question['cohen_kappa']:.3f} | {question['gwet_ac1']:.3f} |",
        f"| Evidence completeness | {completeness['exact_agreement']:.3f} | {completeness['cohen_kappa']:.3f} | {completeness['gwet_ac1']:.3f} |",
        f"| Passage label (5-way) | {five['exact_agreement']:.3f} | {five['cohen_kappa']:.3f} | {five['gwet_ac1']:.3f} |",
        f"| Passage core Gold (binary) | {core['exact_agreement']:.3f} | {core['cohen_kappa']:.3f} | {core['gwet_ac1']:.3f} |",
        "",
        f"Binary core-Gold positive agreement: {core['positive_agreement']:.3f}.",
        f"Macro per-query Gold-set Jaccard: {set_metrics['macro_jaccard']:.3f}; F1: {set_metrics['macro_f1']:.3f}.",
        "",
        "## Items requiring adjudication",
        "",
        f"- Question-level disagreements: {report['counts']['question_disagreements']}",
        f"- Core Gold/non-Gold passage disagreements: {report['counts']['core_passage_disagreements']}",
        f"- Missing passage labels: {report['counts']['missing_passages_a'] + report['counts']['missing_passages_b']}",
        "",
        "Context Only versus Not Supporting disagreements do not alter the core Gold set and are retained for five-way IAA rather than sent to substantive adjudication.",
        "",
    ]
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--reviewer-a", type=Path, required=True)
    parser.add_argument("--reviewer-b", type=Path, nargs=3, required=True)
    parser.add_argument("--registry", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--bootstrap-iterations", type=int, default=10_000)
    args = parser.parse_args()

    registry = read_json(args.registry)
    query_lookup = {row["query_id"]: row for row in registry["queries"]}
    passage_lookup = {
        (query["query_id"], passage["blind_passage_id"]): passage
        for query in registry["queries"] for passage in query["candidate_passages"]
    }
    expected_questions = set(query_lookup)
    expected_passages = set(passage_lookup)
    question_a, passage_a = parse_a_docx(args.reviewer_a)
    question_b, passage_b = parse_b_workbooks(args.reviewer_b)
    issues = validate_labels(question_a, passage_a, "A") + validate_labels(question_b, passage_b, "B")

    if set(question_a) != expected_questions or set(question_b) != expected_questions:
        raise ValueError("question IDs do not match frozen registry")
    extra_passages = (set(passage_a) | set(passage_b)) - expected_passages
    if extra_passages:
        raise ValueError(f"passage IDs absent from frozen registry: {sorted(extra_passages)[:3]}")

    paired_keys = sorted(set(passage_a) & set(passage_b))
    labels_a = [passage_a[key]["passage_label"] for key in paired_keys]
    labels_b = [passage_b[key]["passage_label"] for key in paired_keys]
    question_ids = sorted(expected_questions)
    status_a = [question_a[qid]["question_status"] for qid in question_ids]
    status_b = [question_b[qid]["question_status"] for qid in question_ids]
    completeness_a = [question_a[qid]["evidence_set_completeness"] for qid in question_ids]
    completeness_b = [question_b[qid]["evidence_set_completeness"] for qid in question_ids]
    core_a = [label in CORE_LABELS for label in labels_a]
    core_b = [label in CORE_LABELS for label in labels_b]

    paired_by_query: dict[str, list[tuple[bool, bool]]] = defaultdict(list)
    for key, a_value, b_value in zip(paired_keys, core_a, core_b, strict=True):
        paired_by_query[key[0]].append((a_value, b_value))
    core_metrics = binary_metrics(core_a, core_b)
    core_metrics["cluster_bootstrap_95_ci"] = cluster_bootstrap_core(paired_by_query, args.bootstrap_iterations, SEED)

    set_rows = per_query_set_metrics(passage_a, passage_b, expected_questions)
    valid_jaccard = [row["jaccard"] for row in set_rows if row["jaccard"] is not None]
    valid_f1 = [row["f1"] for row in set_rows if row["f1"] is not None]

    question_disagreements = []
    for query_id in question_ids:
        a_record, b_record = question_a[query_id], question_b[query_id]
        if (a_record["question_status"], a_record["evidence_set_completeness"]) != (b_record["question_status"], b_record["evidence_set_completeness"]):
            query = query_lookup[query_id]
            question_disagreements.append({
                "query_id": query_id,
                "batch": query["batch"],
                "query_slice": query["query_slice"],
                "question": query["query"],
                "a": a_record,
                "b": b_record,
            })

    passage_disagreements = []
    for key in sorted(expected_passages):
        a_record, b_record = passage_a.get(key), passage_b.get(key)
        a_label = "" if a_record is None else a_record["passage_label"]
        b_label = "" if b_record is None else b_record["passage_label"]
        core_disagreement = (a_label in CORE_LABELS) != (b_label in CORE_LABELS)
        missing = not a_label or not b_label
        if core_disagreement or missing:
            query_id, passage_id = key
            query, passage = query_lookup[query_id], passage_lookup[key]
            passage_disagreements.append({
                "query_id": query_id,
                "batch": query["batch"],
                "query_slice": query["query_slice"],
                "question": query["query"],
                "passage_id": passage_id,
                "source_document": passage["source_document"],
                "heading": passage["heading"],
                "frozen_source_passage": passage["content"],
                "a_label": a_label,
                "a_note": "" if a_record is None else a_record["reviewer_note"],
                "b_label": b_label,
                "b_note": "" if b_record is None else b_record["reviewer_note"],
                "reason": "missing_label" if missing else "core_gold_disagreement",
            })

    report = {
        "schema_version": "1.0",
        "status": "initial_independent_labels_analyzed_pending_adjudication",
        "created_at_utc": datetime.now(timezone.utc).isoformat(),
        "bootstrap_iterations": args.bootstrap_iterations,
        "bootstrap_seed": SEED,
        "counts": {
            "questions_a": len(question_a),
            "questions_b": len(question_b),
            "passages_a": len(passage_a),
            "passages_b": len(passage_b),
            "paired_passages": len(paired_keys),
            "missing_passages_a": len(expected_passages - set(passage_a)),
            "missing_passages_b": len(expected_passages - set(passage_b)),
            "question_disagreements": len(question_disagreements),
            "five_way_passage_disagreements": sum(a != b for a, b in zip(labels_a, labels_b, strict=True)),
            "core_passage_disagreements": sum(a != b for a, b in zip(core_a, core_b, strict=True)),
            "queries_with_core_disagreement": len({row["query_id"] for row in passage_disagreements if row["reason"] == "core_gold_disagreement"}),
        },
        "label_validation_issues": issues,
        "agreement": {
            "question_status": agreement_metrics(status_a, status_b, QUESTION_LABELS),
            "evidence_set_completeness": agreement_metrics(completeness_a, completeness_b, COMPLETENESS_LABELS),
            "passage_five_way": agreement_metrics(labels_a, labels_b, PASSAGE_LABELS),
            "passage_core_binary": core_metrics,
        },
        "per_query_gold_set_summary": {
            "queries": len(set_rows),
            "both_empty_queries_excluded": sum(row["jaccard"] is None for row in set_rows),
            "macro_jaccard": statistics.mean(valid_jaccard) if valid_jaccard else None,
            "macro_f1": statistics.mean(valid_f1) if valid_f1 else None,
        },
        "per_query_gold_set_metrics": set_rows,
        "question_disagreements": question_disagreements,
        "passage_disagreements": passage_disagreements,
    }
    args.output_dir.mkdir(parents=True, exist_ok=True)
    write_json(args.output_dir / "dual_annotation_initial_agreement.json", report)
    write_json(args.output_dir / "dual_annotation_disagreement_registry.json", {
        "schema_version": "1.0",
        "status": "pending_joint_adjudication",
        "question_disagreements": question_disagreements,
        "passage_disagreements": passage_disagreements,
    })
    (args.output_dir / "dual_annotation_initial_agreement.md").write_text(markdown_report(report), encoding="utf-8")
    print(json.dumps({"counts": report["counts"], "agreement": report["agreement"], "output_dir": str(args.output_dir)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
