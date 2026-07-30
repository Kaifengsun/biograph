"""Build and format the project-group Word manuscript from the LaTeX source."""

from __future__ import annotations

import argparse
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TITLE = (
    "Text-First Hierarchical Retrieval for Auditable Pharmaceutical "
    "Regulatory Evidence and Supply-Chain Fact Verification"
)

FIGURE_CAPTIONS = {
    "Overview of the proposed dual-path auditable evidence retrieval framework.": 1,
    "Summary of source-chunk and evidence-chain retrieval results.": 2,
}

TABLE_CAPTIONS = {
    "Frozen corpus and evidence-graph statistics.": 1,
    "Independent enrichment ablation": 2,
    "Source-chunk retrieval on 58 adjudicated questions.": 3,
    "Exact evidence-chain ranking on 28 audit-qualified questions.": 4,
}

HEADING_NUMBERS = {
    "Introduction": "1",
    "Related Work": "2",
    "Task Definition": "3",
    "Corpus and Evidence Graph": "4",
    "Methods": "5",
    "Bottom-up source-chunk retrieval": "5.1",
    "Top-down document routing": "5.2",
    "Bounded graph traversal": "5.3",
    "Relation-aware evidence-chain ranking": "5.4",
    "Fusion and adaptive variants": "5.5",
    "BM25-anchored selective reranking": "5.6",
    "Experimental Design": "6",
    "Development, holdout, and confirmation sets": "6.1",
    "Independent annotation and adjudication": "6.2",
    "Methods and metrics": "6.3",
    "Exploratory graph-chain evaluation": "6.4",
    "Results": "7",
    "Historical holdout and initial graph feasibility": "7.1",
    "Enrichment ablation": "7.2",
    "Adjudicated evaluation": "7.3",
    "Neural boundary experiments": "7.4",
    "Exploratory evidence-chain ranking": "7.5",
    "Discussion": "8",
    "A text-first division of labor": "8.1",
    "What the graph contributes": "8.2",
    "Implications for regulatory retrieval evaluation": "8.3",
    "Limitations": "8.4",
    "Conclusion": "9",
}

CROSS_REFERENCE_REPAIRS = {
    "explicit. summarizes how an analyst request":
        "explicit. Figure 1 summarizes how an analyst request",
    "table records. summarizes the resources.":
        "table records. Table 1 summarizes the resources.",
    "nDCG@5 of 0.662 ().":
        "nDCG@5 of 0.662 (Table 2).",
    "MRR of 0.663 ().":
        "MRR of 0.663 (Table 4).",
    "consolidates the principal text-retrieval boundary":
        "Figure 2 consolidates the principal text-retrieval boundary",
}

REFERENCE_TEXTS = (
    "[1] U.S. Food and Drug Administration Drug Shortages Task Force. Drug shortages: Root causes and potential solutions. Technical report, U.S. Food and Drug Administration, 2019. Updated February 2020.",
    "[2] International Council for Harmonisation. ICH Q9(R1): Quality risk management. Technical report, International Council for Harmonisation, 2023.",
    "[3] openFDA. Drug shortages API. U.S. Food and Drug Administration, 2026. Accessed 2026-07-16.",
    "[4] U.S. Food and Drug Administration. Quality systems approach to pharmaceutical current good manufacturing practice regulations: Guidance for industry. Technical report, U.S. Food and Drug Administration, 2006.",
    "[5] European Commission. Eudralex volume 4, annex 11: Computerised systems. Technical report, European Commission, 2011.",
    "[6] Stephen Robertson and Hugo Zaragoza. The probabilistic relevance framework: Bm25 and beyond. Foundations and Trends in Information Retrieval, 4(1–2):1–174, 2009.",
    "[7] Vladimir Karpukhin, Barlas Oguz, Sewon Min, Patrick Lewis, Ledell Wu, Sergey Edunov, Danqi Chen, and Wen-tau Yih. Dense passage retrieval for open-domain question answering. In Proceedings of the 2020 Conference on Empirical Methods in Natural Language Processing, pages 6769–6781, 2020.",
    "[8] Luyu Gao, Xueguang Ma, Jimmy Lin, and Jamie Callan. Precise zero-shot dense retrieval without relevance labels. In Proceedings of the 61st Annual Meeting of the Association for Computational Linguistics, pages 1762–1777, 2023.",
    "[9] Yanzhao Zhang, Mingxin Li, Dingkun Long, Xin Zhang, Huan Lin, Baosong Yang, Pengjun Xie, An Yang, Dayiheng Liu, Junyang Lin, Fei Huang, and Jingren Zhou. Qwen3 embedding: Advancing text embedding and reranking through foundation models. arXiv preprint arXiv:2506.05176, 2025.",
    "[10] Qiao Jin, Won Kim, Qingyu Chen, Donald C. Comeau, Lana Yeganova, W. John Wilbur, and Zhiyong Lu. MedCPT: Contrastive pre-trained transformers with large-scale PubMed search logs for zero-shot biomedical information retrieval. Bioinformatics, 39(11):btad651, 2023.",
    "[11] Parth Sarthi, Salman Abdullah, Aditi Tuli, Shubh Khanna, Anna Goldie, and Christopher D. Manning. RAPTOR: Recursive abstractive processing for tree-organized retrieval. In International Conference on Learning Representations, 2024.",
    "[12] Aidan Hogan, Eva Blomqvist, Michael Cochez, Claudia d’Amato, Gerard de Melo, Claudio Gutierrez, Sabrina Kirrane, Jose Emilio Labra Gayo, Roberto Navigli, Sebastian Neumaier, Axel-Cyrille Ngonga Ngomo, Axel Polleres, Sabbir M. Rashid, Anisa Rula, Lukas Schmelzeisen, Juan Sequeda, Steffen Staab, and Antoine Zimmermann. Knowledge graphs. ACM Computing Surveys, 54(4):71:1–71:37, 2021.",
    "[13] Kalervo Jarvelin and Jaana Kekalainen. Cumulated gain-based evaluation of IR techniques. ACM Transactions on Information Systems, 20(4):422–446, 2002.",
    "[14] Jacob Cohen. A coefficient of agreement for nominal scales. Educational and Psychological Measurement, 20(1):37–46, 1960.",
    "[15] Kilem Li Gwet. Computing inter-rater reliability and its variance in the presence of high agreement. British Journal of Mathematical and Statistical Psychology, 61(1):29–48, 2008.",
)


def set_run_font(run, name: str, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_width(cell, width) -> None:
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_cell_margins(cell, top=60, start=80, bottom=60, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))
    set_run_font(run, "Times New Roman", 9)


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman", 10.5)


def replace_text_preserving_xml(paragraph, old: str, new: str) -> bool:
    nodes = paragraph._p.xpath(".//w:t")
    values = [node.text or "" for node in nodes]
    joined = "".join(values)
    start = joined.find(old)
    if start < 0:
        return False
    end = start + len(old)
    offsets = []
    cursor = 0
    for value in values:
        offsets.append((cursor, cursor + len(value)))
        cursor += len(value)
    first = next(i for i, (_, right) in enumerate(offsets) if right > start)
    last = next(i for i, (_, right) in enumerate(offsets) if right >= end)
    first_left, _ = offsets[first]
    last_left, _ = offsets[last]
    prefix = values[first][: start - first_left]
    suffix = values[last][end - last_left :]
    if first == last:
        nodes[first].text = prefix + new + suffix
    else:
        nodes[first].text = prefix + new
        for index in range(first + 1, last):
            nodes[index].text = ""
        nodes[last].text = suffix
    return True


def prepend_paragraph_text(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman", 9)
    paragraph._p.remove(run._r)
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, run._r)


def make_word_run(
    text: str,
    *,
    vertical_align: str | None = None,
    italic: bool = False,
) -> OxmlElement:
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    properties.append(fonts)
    if italic:
        properties.append(OxmlElement("w:i"))
    if vertical_align is not None:
        align = OxmlElement("w:vertAlign")
        align.set(qn("w:val"), vertical_align)
        properties.append(align)
    run.append(properties)
    node = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        node.set(qn("xml:space"), "preserve")
    node.text = text
    run.append(node)
    return run


def replace_fragile_math_objects(document: Document) -> None:
    math_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    replacements = {"equals": 0, "cq_star": 0, "stars": 0}
    for math in list(document.element.body.iter(f"{{{math_ns}}}oMath")):
        flat = "".join(
            node.text or "" for node in math.iter(f"{{{math_ns}}}t")
        ).replace("\u200b", "")
        parent = math.getparent()
        index = parent.index(math)
        if flat == "=0.918":
            parent.insert(index, make_word_run(" = 0.918"))
            parent.remove(math)
            replacements["equals"] += 1
        elif flat == "Cq*":
            parent.insert(index, make_word_run("C", italic=True))
            parent.insert(index + 1, make_word_run("q", vertical_align="subscript", italic=True))
            parent.insert(index + 2, make_word_run("*", vertical_align="superscript"))
            parent.remove(math)
            replacements["cq_star"] += 1
        elif flat == "*":
            parent.insert(index, make_word_run("*", vertical_align="superscript"))
            parent.remove(math)
            replacements["stars"] += 1
    if replacements != {"equals": 2, "cq_star": 1, "stars": 4}:
        raise RuntimeError(f"Unexpected fragile-math replacements: {replacements}")


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(4)

    for style_name in ("Body Text", "First Paragraph", "Abstract", "Author", "Bibliography"):
        try:
            style = styles[style_name]
        except KeyError:
            continue
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor(0, 0, 0)

    for style_name, size in (("Title", 16), ("Heading 1", 13.5), ("Heading 2", 11.5)):
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if style_name == "Heading 1" else 7)
        style.paragraph_format.space_after = Pt(4)

    if "Manuscript Caption" not in styles:
        caption = styles.add_style("Manuscript Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Manuscript Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = True

    if "Abstract Heading" not in styles:
        abstract_heading = styles.add_style("Abstract Heading", WD_STYLE_TYPE.PARAGRAPH)
    else:
        abstract_heading = styles["Abstract Heading"]
    abstract_heading.font.name = "Times New Roman"
    abstract_heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    abstract_heading.font.size = Pt(11)
    abstract_heading.font.bold = True
    abstract_heading.font.color.rgb = RGBColor(0, 0, 0)
    abstract_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_heading.paragraph_format.space_before = Pt(3)
    abstract_heading.paragraph_format.space_after = Pt(3)
    abstract_heading.paragraph_format.keep_with_next = True


def configure_sections(document: Document) -> None:
    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)
        section.header_distance = Cm(1.1)
        section.footer_distance = Cm(1.1)
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        add_page_number(paragraph)


def repair_cross_references(document: Document) -> None:
    for paragraph in document.paragraphs:
        for old, new in CROSS_REFERENCE_REPAIRS.items():
            replace_text_preserving_xml(paragraph, old, new)


def format_front_matter(document: Document) -> None:
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text == TITLE:
            paragraph.style = document.styles["Title"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(8)
        elif text.startswith("Kaifeng Sun") and "China Jiliang University" in text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(9)
            for run in paragraph.runs:
                set_run_font(run, "Times New Roman", 10.5)
        elif text.startswith("Pharmaceutical analysts must recover"):
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(7)
        elif text.startswith("Keywords:"):
            paragraph.paragraph_format.space_after = Pt(8)
            if paragraph.runs:
                paragraph.runs[0].bold = True
        if index > 10:
            break


def ensure_abstract_heading(document: Document) -> None:
    if any(p.text.strip() == "Abstract" for p in document.paragraphs[:8]):
        return
    abstract = next((p for p in document.paragraphs if p.style.name == "Abstract"), None)
    if abstract is not None:
        heading = abstract.insert_paragraph_before("Abstract")
        heading.style = document.styles["Abstract Heading"]


def restore_heading_numbers(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        number = HEADING_NUMBERS.get(text)
        if number is not None:
            set_paragraph_text(paragraph, f"{number} {text}")
            top_level = "." not in number
            paragraph.style = document.styles["Heading 1" if top_level else "Heading 2"]
            for run in paragraph.runs:
                set_run_font(run, "Times New Roman", 13.5 if top_level else 11.5, True)


def format_captions(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        number = next((n for prefix, n in FIGURE_CAPTIONS.items() if text.startswith(prefix)), None)
        kind = "Figure"
        if number is None:
            number = next((n for prefix, n in TABLE_CAPTIONS.items() if text.startswith(prefix)), None)
            kind = "Table"
        if number is None:
            continue
        if not text.startswith(f"{kind} {number}."):
            prepend_paragraph_text(paragraph, f"{kind} {number}: ")
        paragraph.style = document.styles["Manuscript Caption"]
        paragraph.paragraph_format.keep_with_next = kind == "Table"


def format_figures(document: Document) -> None:
    max_width = Cm(16.35)
    for shape in document.inline_shapes:
        if shape.width > max_width:
            ratio = shape.height / shape.width
            shape.width = max_width
            shape.height = int(max_width * ratio)
    for paragraph in document.paragraphs:
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(2)


def format_tables(document: Document) -> None:
    width_sets = (
        (Cm(13.9), Cm(2.4)),
        (Cm(6.1), Cm(3.4), Cm(3.4), Cm(3.4)),
        (Cm(7.0), Cm(3.1), Cm(3.1), Cm(3.1)),
        (Cm(5.3), Cm(2.75), Cm(2.75), Cm(2.75), Cm(2.75)),
    )
    for table_index, table in enumerate(document.tables):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        widths = width_sets[table_index] if table_index < len(width_sets) else None
        if widths:
            grid = table._tbl.tblGrid
            for child in list(grid):
                grid.remove(child)
            for width in widths:
                grid_col = OxmlElement("w:gridCol")
                grid_col.set(qn("w:w"), str(int(width.twips)))
                grid.append(grid_col)
        if table.rows:
            set_repeat_table_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_table_cell_margins(cell)
                if widths and column_index < len(widths):
                    set_cell_width(cell, widths[column_index])
                for paragraph in cell.paragraphs:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        set_run_font(run, "Times New Roman", 8.3, bold=True if row_index == 0 else None)


def repair_table_data(document: Document) -> None:
    if len(document.tables) < 2:
        return
    rows = document.tables[1].rows
    point_estimates = (
        (2, "0.733", "0.557", "0.526"),
        (3, "0.733", "0.554", "0.524"),
        (4, "0.567", "0.384", "0.357"),
        (5, "0.567", "0.384", "0.357"),
        (6, "0.733", "0.606", "0.508"),
        (7, "0.800", "0.608", "0.571"),
    )
    for row_index, *values in point_estimates:
        for column_index, value in enumerate(values, start=1):
            cell = rows[row_index].cells[column_index]
            current = cell.text.strip()
            if not current.startswith(value):
                cell.text = f"{value} {current}".strip()


def ensure_references_heading(document: Document) -> None:
    if any(p.text.strip() == "References" for p in document.paragraphs):
        return
    first_reference = next(
        (p for p in document.paragraphs if p.style.name == "Bibliography"),
        None,
    )
    if first_reference is not None:
        heading = first_reference.insert_paragraph_before("References")
        heading.style = document.styles["Heading 1"]


def restore_pdf_references(document: Document) -> None:
    bibliography = [p for p in document.paragraphs if p.style.name == "Bibliography"]
    if len(bibliography) != len(REFERENCE_TEXTS):
        raise RuntimeError(
            f"Expected {len(REFERENCE_TEXTS)} bibliography entries, found {len(bibliography)}"
        )
    for paragraph, reference in zip(bibliography, REFERENCE_TEXTS):
        set_paragraph_text(paragraph, reference)
        paragraph.style = document.styles["Bibliography"]


def number_native_display_equations(document: Document) -> None:
    math_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    display_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph._p.find(f"{{{math_ns}}}oMathPara") is not None
    ]
    if len(display_paragraphs) != 4:
        raise RuntimeError(
            f"Expected 4 display equations, found {len(display_paragraphs)}"
        )
    for number, paragraph in enumerate(display_paragraphs, start=1):
        math_para = paragraph._p.find(f"{{{math_ns}}}oMathPara")
        math = math_para.find(f"{{{math_ns}}}oMath")
        math_para.remove(math)
        paragraph._p.remove(math_para)

        p_pr = paragraph._p.get_or_add_pPr()
        tabs = p_pr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            p_pr.append(tabs)
        for child in list(tabs):
            tabs.remove(child)
        center_tab = OxmlElement("w:tab")
        center_tab.set(qn("w:val"), "center")
        center_tab.set(qn("w:pos"), "4680")
        right_tab = OxmlElement("w:tab")
        right_tab.set(qn("w:val"), "right")
        right_tab.set(qn("w:pos"), "9360")
        tabs.extend((center_tab, right_tab))

        leading_run = paragraph.add_run()
        leading_run._r.append(OxmlElement("w:tab"))
        paragraph._p.append(math)
        label_run = paragraph.add_run()
        label_run._r.append(OxmlElement("w:tab"))
        label_run.add_text(f"({number})")
        set_run_font(label_run, "Times New Roman", 10.5)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)


def match_pdf_float_order(document: Document) -> None:
    heading = next(
        p for p in document.paragraphs if p.text.strip() == "7.4 Neural boundary experiments"
    )
    caption = next(
        p for p in document.paragraphs if p.text.startswith("Table 3:")
    )
    note = next(
        p
        for p in document.paragraphs
        if "Feedback-motivated locked extensions on already observed Gold questions" in p.text
    )
    table = document.tables[2]
    heading._p.addprevious(caption._p)
    heading._p.addprevious(table._tbl)
    heading._p.addprevious(note._p)


def format_remaining_paragraphs(document: Document) -> None:
    in_references = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text == "References":
            in_references = True
        if paragraph.style.name not in {
            "Title",
            "Author",
            "Heading 1",
            "Heading 2",
            "Abstract Heading",
            "Manuscript Caption",
        }:
            if not paragraph._p.xpath(".//w:drawing"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.paragraph_format.space_after = Pt(4)
        if in_references and text != "References":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Cm(0.55)
            paragraph.paragraph_format.first_line_indent = Cm(-0.55)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                set_run_font(run, "Times New Roman", 8.7)
        else:
            for run in paragraph.runs:
                if run.font.name is None:
                    set_run_font(run, "Times New Roman", 10.5)
        if text.startswith((
            "Counts come from the frozen enrichment",
            "R1–R4 are cumulative flat dense indexes",
            "*Feedback-motivated locked extensions",
            "Questions were constructed from known graph relations.",
        )):
            paragraph.paragraph_format.left_indent = Cm(0.25)
            paragraph.paragraph_format.right_indent = Cm(0.25)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(5)
            for run in paragraph.runs:
                set_run_font(run, "Times New Roman", 8.3)
                run.italic = True


def build_document(
    source: Path,
    output: Path,
    pandoc: Path,
    bibliography: Path,
    native_math: bool = False,
) -> None:
    script_dir = Path(__file__).resolve().parent
    with tempfile.TemporaryDirectory(prefix="word_manuscript_") as temp_dir:
        intermediate = Path(temp_dir) / "intermediate.docx"
        command = [
            str(pandoc),
            str(source),
            "--from=latex",
            "--to=docx",
            f"--resource-path={source.parent};{source.parent / 'figures'}",
            f"--bibliography={bibliography}",
            f"--csl={script_dir / 'word_numeric.csl'}",
            "--citeproc",
            f"--output={intermediate}",
        ]
        if not native_math:
            command.insert(-2, f"--lua-filter={script_dir / 'plain_math.lua'}")
        subprocess.run(command, cwd=source.parent, check=True)

        document = Document(intermediate)
        document.core_properties.title = TITLE
        document.core_properties.author = "Kaifeng Sun"
        document.core_properties.subject = "Auditable pharmaceutical regulatory evidence retrieval"
        configure_styles(document)
        configure_sections(document)
        if native_math:
            replace_fragile_math_objects(document)
        repair_cross_references(document)
        repair_table_data(document)
        ensure_abstract_heading(document)
        restore_heading_numbers(document)
        format_front_matter(document)
        format_captions(document)
        format_figures(document)
        format_tables(document)
        restore_pdf_references(document)
        ensure_references_heading(document)
        format_remaining_paragraphs(document)
        if native_math:
            number_native_display_equations(document)
        match_pdf_float_order(document)
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=Path("main.tex"))
    parser.add_argument("--output", type=Path, default=Path("skf_manuscript.docx"))
    parser.add_argument("--pandoc", type=Path, default=Path(r"D:\Anaconda3\Scripts\pandoc.exe"))
    parser.add_argument("--bibliography", type=Path, default=Path("references.bib"))
    parser.add_argument("--native-math", action="store_true")
    args = parser.parse_args()
    source = args.source.resolve()
    output = args.output.resolve()
    bibliography = args.bibliography.resolve()
    build_document(source, output, args.pandoc.resolve(), bibliography, args.native_math)
    print(output)


if __name__ == "__main__":
    main()
